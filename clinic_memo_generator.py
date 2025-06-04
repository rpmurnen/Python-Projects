import os
import time
from datetime import datetime
import calendar
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


try:
    import xlwings as xw
except ImportError:
    raise ImportError("The 'xlwings' module is required. Please install it using: pip install xlwings")

wb = None  # Predefine workbook variable so it always exists
prior_app = None
prior_wb = None

# === CONFIGURATION ===
EXCEL_PATH = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\05-2025\05-25 CLINIC FINANCIALS.xlsm")
PRIOR_EXCEL_PATH = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\04-2025\04-25 CLINIC FINANCIALS2.xlsm")
DEST_FOLDER_BASE = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\05-2025\Clinics\MEMO COVER PAGE")

WORD_TEMPLATE_PATH = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\ME_Reporting_Python_Scripting\memo_template.docx")
TODAY_FOLDER = datetime.today().strftime("%Y-%m-%d")
DEST_FOLDER = os.path.join(DEST_FOLDER_BASE, TODAY_FOLDER)
os.makedirs(DEST_FOLDER, exist_ok=True)

CONTROL_SHEET = "Control"
AR_SHEET = "AR_Patient_Report"
CLINIC_LIST_RANGE = "B6:B34"
AR_DROPDOWN_CELL = "B2"
PL_DROPDOWN_CELL = 'D3'
BS_DROPDOWN_CELL = 'E3'
TABLE1_RANGE = "C11:G15"
TABLE2_RANGE = "C18:G28"

PLACEHOLDER_1 = "<<INSERT EXCEL SECTION 1 HERE>>"
PLACEHOLDER_2 = "<<INSERT EXCEL SECTION 2 HERE>>"

# Add your properly indented and corrected code logic below this line
# Let me know if you want the full corrected script inserted here
import os
import time
from datetime import datetime
import calendar
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

try:
    import xlwings as xw
except ImportError:
    raise ImportError("The 'xlwings' module is required. Please install it using: pip install xlwings")

wb = None  # Predefine workbook variable so it always exists
prior_app = None
prior_wb = None

# === CONFIGURATION ===
# === YOU NEED TO DEFINE THESE VARIABLES:===
EXCEL_PATH = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\05-2025\05-25 CLINIC FINANCIALS.xlsm")
PRIOR_EXCEL_PATH = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\04-2025\04-25 CLINIC FINANCIALS2.xlsm")
DEST_FOLDER_BASE = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\05-2025\Clinics\MEMO COVER PAGE")

# === DON'T CHANGE THIS ONE - IT'S THE MEMO TEMPLATE THE SCRIPT WILL USE:===
WORD_TEMPLATE_PATH = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\ME_Reporting_Python_Scripting\memo_template.docx")

# === THIS IS THE SCRIPT SETTING TODAY'S TIME AND DATE....THE SCRIPT MUST BE RUN IN MONTH:===
TODAY_FOLDER = datetime.today().strftime("%Y-%m-%d")
DEST_FOLDER = os.path.join(DEST_FOLDER_BASE, TODAY_FOLDER)
os.makedirs(DEST_FOLDER, exist_ok=True)

CONTROL_SHEET = "Control"
AR_SHEET = "AR_Patient_Report"
CLINIC_LIST_RANGE = "B6:B34"
AR_DROPDOWN_CELL = "B2"
PL_DROPDOWN_CELL = 'D3'
BS_DROPDOWN_CELL = 'E3'
TABLE1_RANGE = "C11:G15"
TABLE2_RANGE = "C18:G28"

PLACEHOLDER_1 = "<<INSERT EXCEL SECTION 1 HERE>>"
PLACEHOLDER_2 = "<<INSERT EXCEL SECTION 2 HERE>>"




# === HELPER FUNCTIONS:===
# get patient counts
def get_patient_count(ws):
    try:
        val = ws.range("G15").value
        return int(val) if val is not None else 0
    except:
        return 0


# get the patient counts by doctor
def get_doctor_patient_stats(ws, row):
    name = ws.range(f"C{row}").value or "[Name missing]"
    primaries = ws.range(f"D{row}").value or 0
    spouse = ws.range(f"E{row}").value or 0
    children = ws.range(f"F{row}").value or 0
    total = ws.range(f"G{row}").value or 0
    return {
        "name": str(name),
        "primaries": int(primaries),
        "spouse": int(spouse),
        "children": int(children),
        "total": int(total)
    }



def get_doctor_names_xlwings(xlwb, clinic_name):
    md_list_sheet = xlwb.sheets["MD_List"]
    data = md_list_sheet.range("A2:F100").value  # Adjust if needed

    for row in data:
        if row[0] and str(row[0]).strip().lower() == clinic_name.strip().lower():
            # Columns C, D = full names; Columns E, F = last names
            full_names = [str(row[2]).strip(), str(row[3]).strip()]
            last_names = [str(row[4]).strip(), str(row[5]).strip()]
            full_names = [name for name in full_names if name]
            last_names = [name for name in last_names if name]
            return full_names, last_names

    return [], []


def get_previous_month_year(month_offset=1):
    today = datetime.today()
    month = today.month - month_offset
    year = today.year
    if month <= 0:
        month += 12
        year -= 1
    month_name = calendar.month_name[month]
    return month_name, year

def get_cash_balance(bs_sheet, target_month):
    # Find the last used column from row 4
    header = bs_sheet.range("A4:Z4").value
    if not header:
        return None

    for idx, month_label in enumerate(header):
        if month_label and str(month_label).strip().lower() == target_month.lower():
            # Convert index to Excel column letter (A=1)
            col_letter = chr(ord('A') + idx)
            cell = f"{col_letter}6"
            return bs_sheet.range(cell).value

    return None


def build_doctor_commentary(current, prior):
    name = current["name"]
    total_diff = current["total"] - prior["total"]
    prim_diff = current["primaries"] - prior["primaries"]
    spouse_diff = current["spouse"] - prior["spouse"]
    child_diff = current["children"] - prior["children"]

    def fmt(val):
        sign = "+" if val > 0 else ""  # no sign for negative since "-" will be included
        return f"{sign}{val}"

    return (
        f"{name}: total patient volume change = {fmt(total_diff)} "
        f"(Primaries: {fmt(prim_diff)}, Spouse/Mid-Adult: {fmt(spouse_diff)}, Children: {fmt(child_diff)})"
    )


def set_font(run):
    run.font.name = 'Aptos Narrow'
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos Narrow')


# formatting helper - bullet points
def add_bullet_paragraph(doc, text, index):
    para = doc.paragraphs[index].insert_paragraph_before(text)
    set_font(para.runs[0])

    # Add bullet formatting manually
    p = para._element
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl)
    numPr.append(numId)
    p.insert(0, numPr)

    return para

# ... [imports and config remain unchanged above this point]
# Script begins below this point:


# This section drops in the tables from the financials xlsx - looks for the tag in the memo doc then replaces
def insert_table_at_placeholder(doc, placeholder, table_data):
    print(f"\n🔍 Looking for placeholder: {placeholder}")
    found = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        print(f"👀 Paragraph {i}: '{text}'")

        if placeholder in text:
            print(f"✅ Found placeholder in paragraph {i}")
            para.text = para.text.replace(placeholder, "")
            found = True

            table = doc.add_table(rows=0, cols=len(table_data[0]))
            table.style = 'Table Grid'

            for row_idx, row_data in enumerate(table_data):
                row_cells = table.add_row().cells
                if row_idx == 0 and placeholder == PLACEHOLDER_1:
                    print("🔧 Merging and centering header row (cols 1–4) for Table 1")
                    
                    # Merge columns 1 to 4
                    merged_cell = row_cells[1]
                    for col in range(2, 5):
                        merged_cell = merged_cell.merge(row_cells[col])
                    
                    # Center and style the merged cell
                    merged_paragraph = merged_cell.paragraphs[0]
                    merged_paragraph.alignment = 1  # Center
                    merged_paragraph.clear()
                    run = merged_paragraph.add_run(str(table_data[0][1]) if table_data[0][1] else "")
                    set_font(run)
                    run.bold = True

                # Merge and left-align for Table 2 title row
                if row_idx == 0 and placeholder == PLACEHOLDER_2:
                    print("🔧 Merging and left-aligning title row for Table 2")
                    merged_cell = row_cells[0]
                    for col in range(1, len(row_cells)):
                        merged_cell = merged_cell.merge(row_cells[col])

                    paragraph = merged_cell.paragraphs[0]
                    paragraph.alignment = 0  # Left align
                    paragraph.clear()
                    run = paragraph.add_run(str(table_data[0][0]) if table_data[0][0] else "")
                    set_font(run)
                    run.bold = True

                    continue  # ✅ Skip rest of row so it doesn't get re-added below

                for col_idx, cell_val in enumerate(row_data):
                    if row_idx == 0 and placeholder == PLACEHOLDER_1 and col_idx >= 1:
                        continue  # Skip all merged cells: 1, 2, 3, 4


                    cell = row_cells[col_idx]
                    # Format values
                    if isinstance(cell_val, (int, float)):
                        text = f"{cell_val:,.0f}"
                        is_numeric = True
                    elif isinstance(cell_val, datetime):
                        text = cell_val.strftime("%#m/%#d/%y")
                        is_numeric = False
                    elif cell_val is None:
                        text = ""
                        is_numeric = False
                    else:
                        text = str(cell_val)
                        is_numeric = False

                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(text)
                    set_font(run)
                    if is_numeric:
                        paragraph.alignment = 2
                    if placeholder == PLACEHOLDER_1 and row_idx == 1:
                        run.bold = True
                        paragraph.alignment = 1  # Center


                    if row_idx == 0 or row_idx == len(table_data) - 1:
                        run.bold = True

            tbl_element = table._tbl
            para_element = para._element
            para_element.addnext(tbl_element)

            print(f"✅ Table inserted directly after paragraph {i} for: {placeholder}")
            return

    if not found:
        print(f"⚠️ Placeholder not found: {placeholder}")




prior_app = xw.App(visible=False)
prior_app.display_alerts = False
prior_app.screen_updating = False
prior_wb = prior_app.books.open(PRIOR_EXCEL_PATH)


# =========== generate basic clinic commentary function ==========


def generate_commentary(pl_ws):
    def fetch(row, col):
        val = pl_ws.range(f"{col}{row}").value or 0
        return float(val)

    def fmt(val):
        return f"${val:,.0f}"

    def build_line(label, actual, budget):
        delta = actual - budget
        pct = (delta / budget) if budget else 0
        abs_pct = abs(pct)
        abs_delta = abs(delta)

        if abs_pct < 0.05:
            tone = "in line with plan"
        elif abs_pct < 0.20 or abs_delta < 10000:
            tone = "moderately above plan" if delta > 0 else "moderately below plan"
        else:
            tone = "significantly above plan" if delta > 0 else "significantly below plan"

        actual_str = f"${actual:,.0f}"
        budget_str = f"${budget:,.0f}"
        delta_str = f"${abs_delta:,.0f} ({abs_pct:.0%})"

        return f"{label} finished at {actual_str} versus a budget of {budget_str}, a difference of {delta_str} - {tone}.".strip()

    actuals = {
        "net": fetch(83, "F"),
        "rev": fetch(11, "F"),
        "payroll": fetch(26, "F"),
        "patient": fetch(32, "F"),
        "occupancy": fetch(45, "F"),
        "taxes": fetch(58, "F"),
        "dues": fetch(65, "F"),
    }
    budgets = {
        "net": fetch(83, "G"),
        "rev": fetch(11, "G"),
        "payroll": fetch(26, "G"),
        "patient": fetch(32, "G"),
        "occupancy": fetch(45, "G"),
        "taxes": fetch(58, "G"),
        "dues": fetch(65, "G"),
    }

    
    total_exp_actual = fetch(81, "F")
    total_exp_budget = fetch(81, "G")
    exp_delta = total_exp_budget - total_exp_actual

    rev_delta = actuals["rev"] - budgets["rev"]

    net_actual = actuals["net"]
    net_budget = budgets["net"]
    net_delta = net_actual - net_budget
    net_pct = (net_delta / net_budget) if net_budget else 0
    abs_net_delta = abs(net_delta)

    if abs_net_delta < 0.05 * net_budget:
        tone = "in line with plan"
    elif abs_net_delta < 0.20 * net_budget or abs_net_delta < 10000:
        tone = "moderately above plan" if net_delta > 0 else "moderately below plan"
    else:
        tone = "significantly above plan" if net_delta > 0 else "significantly below plan"

    net_line = (
        f"Net income finished at ${net_actual:,.0f} versus a budget of ${net_budget:,.0f}, "
        f"a difference of ${abs_net_delta:,.0f} ({abs(net_pct):.0%}) - {tone}. "
        f"This result was driven by revenue {'exceeding' if rev_delta > 0 else 'coming in below'} "
        f"plan by ${abs(rev_delta):,.0f} and expenses {'coming in under' if exp_delta > 0 else 'exceeding'} "
        f"budget by ${abs(exp_delta):,.0f}."
    )

    return {
        "<<COMMENTARY_NET_INCOME>>": net_line,
        "<<COMMENTARY_REVENUE>>": build_line("Revenue", actuals["rev"], budgets["rev"]),
        "<<COMMENTARY_PAYROLL>>": build_line("Payroll & Related", actuals["payroll"], budgets["payroll"]),
        "<<COMMENTARY_PATIENT_OP_EXP>>": build_line("Patient Operating Expenses", actuals["patient"], budgets["patient"]),
        "<<COMMENTARY_OCCUPANCY>>": build_line("Occupancy", actuals["occupancy"], budgets["occupancy"]),
        "<<COMMENTARY_TAXES>>": build_line("Taxes and Licenses", actuals["taxes"], budgets["taxes"]),
        "<<COMMENTARY_DUES>>": build_line("Fees & Dues", actuals["dues"], budgets["dues"]),

    }




print("🔄 Launching Excel...")
app = xw.App(visible=False)
app.display_alerts = False
app.screen_updating = False

try:
    print("📂 Opening workbook...")
    wb = app.books.open(EXCEL_PATH)
    control_ws = wb.sheets[CONTROL_SHEET]
    ar_ws = wb.sheets[AR_SHEET]

    print("📋 Fetching clinic list...")
    raw_clinics = control_ws.range(CLINIC_LIST_RANGE).value
    clinic_list = [name for name in raw_clinics if name]

    print(f"✅ Found {len(clinic_list)} clinics.")





    for clinic in clinic_list:
        print(f"\n🩺 Processing clinic: {clinic}")
        bs_ws = wb.sheets["BS"]
        month_name, _ = get_previous_month_year()
        ar_ws.range(BS_DROPDOWN_CELL).value = clinic
        time.sleep(0.25)
        _ = ar_ws.range("C11").value
        ar_ws.api.Calculate()
        pl_ws = wb.sheets["PL"]

        # Set dropdowns to the current clinic for all sheets
        ar_ws.range(AR_DROPDOWN_CELL).value = clinic
        pl_ws.range(PL_DROPDOWN_CELL).value = clinic
        bs_ws.range(BS_DROPDOWN_CELL).value = clinic

        # Force recalc by referencing a known formula-driven cell
        _ = ar_ws.range("C11").value
        ar_ws.api.Calculate()
        _ = pl_ws.range("F83").value
        pl_ws.api.Calculate()
        _ = bs_ws.range("F6").value  # optional, if needed
        bs_ws.api.Calculate()

        cash_balance = get_cash_balance(bs_ws, month_name)





        # Switch clinic in prior workbook too
        prior_ar_ws = prior_wb.sheets[AR_SHEET]
        prior_ar_ws.range(AR_DROPDOWN_CELL).value = clinic
        _ = prior_ar_ws.range("C11").value
        prior_ar_ws.api.Calculate()

        prior_count = get_patient_count(prior_ar_ws)
        current_count = get_patient_count(ar_ws)

        delta = current_count - prior_count
        abs_pct = abs(delta) / prior_count if prior_count else 0
        direction = "increased" if delta > 0 else "decreased"

        if delta == 0:
            patient_msg = f"Patient volume was flat compared to the prior month, which had {prior_count} patients."
        else:
            patient_msg = (
                f"Patient volume {direction} by {abs(delta)} "
                f"({abs_pct:.0%}) compared to last month, which had {prior_count} patients."
    )


        # here we actually call all the patient stat and doctor commentary functions we built above
        doc1_current = get_doctor_patient_stats(ar_ws, 13)
        doc2_current = get_doctor_patient_stats(ar_ws, 14)
        doc1_prior = get_doctor_patient_stats(prior_ar_ws, 13)
        doc2_prior = get_doctor_patient_stats(prior_ar_ws, 14)
        doc1_line = build_doctor_commentary(doc1_current, doc1_prior)
        doc2_line = build_doctor_commentary(doc2_current, doc2_prior)



        # Now generate commentary using the updated PL sheet
        commentary_dict = generate_commentary(pl_ws)

        table1 = ar_ws.range(TABLE1_RANGE).value
        table2 = ar_ws.range(TABLE2_RANGE).value

        print(f"\U0001F4CB Clinic: {clinic}")
        for row in table1:
            print(row)

        # Initialize Word document template    
        doc = Document(WORD_TEMPLATE_PATH)

        # calls the doctor names and writes the header of the memo doc
        full_names, last_names = get_doctor_names_xlwings(wb, clinic)
        greeting_line = "Hello " + " and ".join(f"{name}" for name in last_names) + "," if last_names else "Hello,"
        to_line = "To: " + " and ".join(full_names) if full_names else "To: [Doctor Name]"
        from_line = "From: Russ Murnen"
        month_name, year = get_previous_month_year()
        date_line = datetime.today().strftime("%B %d, %Y")
        re_line = f"Re: {month_name} Financial Performance"
        open_line = f"I’m pleased to share your financial results for {month_name} {year}."

        for text in [open_line, "", greeting_line, "", re_line, from_line, to_line, "", date_line]:
            para = doc.paragraphs[0].insert_paragraph_before("")
            run = para.add_run(text)
            set_font(run)

        insert_table_at_placeholder(doc, PLACEHOLDER_1, table1)

        for placeholder, text in commentary_dict.items():
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, text)
                    set_font(para.runs[0])

        # Count non-empty names in C20:C27 (table2[1:8], column 0)
        name_cells = [row[0] for row in table2[1:8]]
        outstanding_count = sum(1 for name in name_cells if isinstance(name, str) and name.strip())

        if outstanding_count == 1:
            ar_message = "You have 1 patient with an outstanding balance > 30 days."
        else:
            ar_message = f"You have {outstanding_count} patients with an outstanding balance > 30 days."


        # Insert patient volume commentary
        for i, para in enumerate(doc.paragraphs):
            if "<<COMMENTARY_PATIENT_VOLUME>>" in para.text:
                clean_text = para.text.replace("<<COMMENTARY_PATIENT_VOLUME>>", "").strip()
                para.text = clean_text

                # Doctor 2 detail (bullet)
                add_bullet_paragraph(doc, doc2_line, i)

                # Doctor 1 detail (bullet)
                add_bullet_paragraph(doc, doc1_line, i)

                # Overall volume message (regular)
                vol_para = doc.paragraphs[i].insert_paragraph_before(patient_msg)
                set_font(vol_para.runs[0])


                break



        # Insert AR outstanding message
        for i, para in enumerate(doc.paragraphs):
            if PLACEHOLDER_2 in para.text:
                ar_para = doc.paragraphs[i].insert_paragraph_before()
                run = ar_para.add_run(ar_message)
                set_font(run)
                break

        table2 = [row for row in table2 if any(cell not in [None, "", " "] for cell in row)]
        insert_table_at_placeholder(doc, PLACEHOLDER_2, table2)

        if cash_balance is not None:
            cash_msg = f"Your cash balance as of the end of {month_name} was ${cash_balance:,.0f}."
        else:
            cash_msg = f"Cash balance for {month_name} not found."

        para = doc.add_paragraph()
        run = para.add_run(cash_msg)
        set_font(run)

        filename = f"{clinic} review.docx"
        output_path = os.path.join(DEST_FOLDER, filename)
        doc.save(output_path)
        print(f"📏 Saved: {output_path}")

except Exception as e:
    print(f"❌ Error occurred: {e}")

finally:
    if wb:
        wb.close()
    if prior_wb:
        prior_wb.close()
    if app:
        app.quit()
    if prior_app:
        prior_app.quit()
