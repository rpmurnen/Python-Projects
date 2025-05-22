import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image
from PyPDF2 import PdfMerger
import win32com.client as win32

# === Folders ===
memo_folder = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\04-2025\Clinics\MEMO COVER PAGE\2025-05-22")
image_folder = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\04-2025\Clinics\Clinic_ScreenShots")
accounting_pdf_folder = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\04-2025\Clinics\PDF TRANSACTION DETAIL")
output_folder = os.path.expanduser(r"~\MD2\Finance - General\00 Financials\2025\04-2025\Clinics\PDF ENTIRE PACKET")
os.makedirs(output_folder, exist_ok=True)

# === Temp folder for intermediate PDFs ===
temp_pdf_folder = os.path.join(output_folder, "temp_parts")
os.makedirs(temp_pdf_folder, exist_ok=True)

# === Get clinic names from memo files ===
clinic_names = []
for file in os.listdir(memo_folder):
    if file.lower().endswith(" review.docx"):
        name = re.sub(r" review\.docx$", "", file, flags=re.IGNORECASE).strip()
        clinic_names.append(name)

print(f"📋 Found {len(clinic_names)} clinics")

# === Convert .docx to PDF ===
def convert_docx_to_pdf(input_path, output_path):
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()

# === Convert image file to a PIL-safe image if needed ===
def ensure_rgb_image(img_path):
    img = Image.open(img_path)
    if img.mode == "RGBA":
        img = img.convert("RGB")
    return img

# === Main loop ===
for clinic in clinic_names:
    print(f"\n📦 Assembling: {clinic}")

    try:
        base_name = clinic.strip()

        # Paths
        memo_docx_path = os.path.join(memo_folder, f"{base_name} Review.docx")
        pl_img_path = os.path.join(image_folder, f"{base_name}_PL.png")
        bs_img_path = os.path.join(image_folder, f"{base_name}_BS.png")
        accounting_pdf_path = os.path.join(accounting_pdf_folder, f"{base_name}.pdf")
        combined_docx_path = os.path.join(temp_pdf_folder, f"{base_name}_combined.docx")
        memo_pdf_path = os.path.join(temp_pdf_folder, f"{base_name}_memo.pdf")
        final_pdf_path = os.path.join(output_folder, f"{base_name} - packet.pdf")

        # 1. Load memo doc
        doc = Document(memo_docx_path)

        # 2. Add PL image
        doc.add_page_break()
        doc.add_picture(pl_img_path, width=Inches(6.5))

        # 3. Add BS image
        doc.add_page_break()
        doc.add_picture(bs_img_path, width=Inches(6.5))

        # 4. Save updated Word file
        doc.save(combined_docx_path)

        # 5. Convert updated Word doc to PDF
        convert_docx_to_pdf(combined_docx_path, memo_pdf_path)

        # 6. Merge memo PDF + accounting PDF
        merger = PdfMerger()
        merger.append(memo_pdf_path)

        if os.path.exists(accounting_pdf_path):
            merger.append(accounting_pdf_path)
        else:
            print(f"⚠️ Missing accounting PDF: {accounting_pdf_path}")

        merger.write(final_pdf_path)
        merger.close()

        print(f"✅ Saved packet: {final_pdf_path}")

    except Exception as e:
        print(f"❌ Error for {clinic}: {e}")

print("\n🎉 All packets generated with memo + images + detail PDFs.")
